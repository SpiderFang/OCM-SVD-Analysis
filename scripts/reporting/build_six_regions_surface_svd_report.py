#!/usr/bin/env python3
"""建立指定海域 2024–2025 表層多變量 SVD 學術成果報告。

本模組只讀取已發布且不可變更的 SVD 科學陣列，以及
``academic_report_ready_v8`` 中已附向量比例尺的 PNG 圖檔；它不會重新讀取 OCM
surface cache、不會重新插補，也不會重新求解 SVD。輸入資料的物理意義、缺值處理、
矩陣建構、尺度正規化與面積加權均於報告方法章明示，並附可重跑的設定檔與命令，讓
讀者能從相同 cache 重現科學 run，再由既有陣列重繪完全相同的 v8 圖集。

輸出為 Microsoft Word ``.docx``。報告內所有地圖均使用
``*_with_vector_scale.png``，確保每一幅平均場與空間模態圖都保留明確的向量比例尺；
PC 與解釋變異量則採同一 v8 bundle 的報告版 PNG。限制是本程式僅負責彙整、計算報告
所需的摘要統計與排版，不改變來源科學結論。聚焦版以各研究海域獨立呈現流場模態、
PC 時序變化及主導海廢輸送候選模態，不進行跨區排序或關聯推論。
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

import numpy as np
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
RESULT_ROOT = PROJECT_ROOT / "work/server_results/2024_2025"
OUTPUT_DEFAULT = PROJECT_ROOT / "outputs/reports/指定海域_流場模態萃取時序變化與主導海廢輸送候選模態_聚焦版_v2_2024-2025.docx"

# 報告版面採 documents skill 的 narrative_proposal 基準；中文與拉丁字元均依委託格式
# 指定為「標楷體-繁」，避免中英文混排時由 Office 自動切換到另一套字族。
# LibreOffice 在 macOS 對 ``w:eastAsia`` fallback 與 variable font 的處理不完全一致；
# 依研究報告格式要求，正文、標題、表格與圖說均指定 macOS／Word 顯示名稱
#「標楷體-繁」。Latin 字元亦使用同一字族，避免中英文混排時字形突然切換；數學公式
# 仍單獨使用 Cambria Math，以確保希臘字母、矩陣符號與上下標完整。若執行環境未安裝
# 此字型，Office 會依系統 fallback 顯示，但 DOCX 內保存的指定字型名稱不會被改寫。
LATIN_FONT = "標楷體-繁"
CJK_FONT = "標楷體-繁"
# 委託格式要求所有可編輯文字都使用同一字族，因此公式也直接指定「標楷體-繁」。若公式
# 所需符號不在該字型內，Office 可能只對缺少的單一符號套用系統 fallback，但 DOCX 的字型
# 指定仍維持一致，不再主動切換至 Cambria Math。
MATH_FONT = "標楷體-繁"
NAVY = "17365D"
BLUE = "1F4E79"
PALE_BLUE = "D9EAF7"
PALE_GRAY = "F2F4F5"
MID_GRAY = "66737F"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360


@dataclass(frozen=True)
class RegionSpec:
    """描述單一分析區的固定名稱、來源 run 與科學判讀。

    ``mode_texts`` 依序對應前五模態的空間結構解讀。這些文字只描述 PC 為正時的回歸場；
    PC 為負時，流向與海表面高度異常符號必須整體反轉。``transport_choice`` 則把「聯合
    變異量最大」與「最可能主導海廢速度擾動」分開陳述，使各海域的候選模態判讀維持
    在自身流場與 PC 時序脈絡內。
    """

    key: str
    name_zh: str
    run_id: str
    location_note: str
    status_zh: str
    transport_choice: str
    overall_interpretation: str
    mode_texts: tuple[str, str, str, str, str]


REGIONS: tuple[RegionSpec, ...] = (
    RegionSpec(
        key="beigan",
        name_zh="北竿",
        run_id="beigan_surface_u_v_eta_available_2024_2025_v1_aa8e78ce1497",
        location_note="馬祖北側島嶼群及鄰近水道",
        status_zh="已核定分析區（analysis_ready）",
        transport_choice="Mode 1 為主要候選；其東至東北向速度異常最強且空間同向性高。Mode 2 主要表現海面高度共同變化，Mode 3 為次要南向擾動。",
        overall_interpretation="前兩模態已高度濃縮區內共同變異；Mode 1 表現東至東北向廣域輸送擾動，Mode 2 則以海面高度共同升降及島間調整流為主。",
        mode_texts=(
            "正相位呈全域一致的東至東北向流速異常，海面高度異常則大致為負；速度回歸場具高度同向性，表示此模態可有效改變島嶼兩側與水道內的水平輸送方向。",
            "海面高度呈全域正異常，水平流場則具有島間轉向與剪切結構；聯合變異量雖高，但速度異常相較 Mode 1 較弱，因此較適合解讀為水位共同升降伴隨局地調整流。",
            "流速異常以南至西南向為主，海面高度多為負；其解釋變異量已顯著低於前兩模態，代表較次要、但可能在特定時段改變南向輸送的擾動。",
            "速度與海面高度回歸場均較局地化，島嶼附近出現相反方向的補償結構；其區域整體輸送代表性低於前三模態。",
            "主要表現為低能量、局部反向的細尺度結構，可能反映高階剪切、殘餘潮汐或取樣噪聲；不列為本區主導海廢輸送候選。",
        ),
    ),
    RegionSpec(
        key="gongliao",
        name_zh="貢寮",
        run_id="gongliao_surface_u_v_eta_available_2024_2025_v1_056dc6b85b34",
        location_note="臺灣東北角近岸及外海",
        status_zh="候選試驗區（candidate_pilot）",
        transport_choice="Mode 1 為主要候選；其東南向速度異常強、同向性高且解釋變異量最大。Mode 2 為具季節相位的南向次要模態。",
        overall_interpretation="Mode 1 控制最廣域的東南向異常輸送；Mode 2 在區域東半部具有較強南向訊號，且月平均 PC 顯示明顯季節差異，可能調節東北角近岸與外海交換。",
        mode_texts=(
            "正相位為全區一致的東南向流速異常，海面高度多為負；速度幅度與空間同向性均為前五模態最高，故同時符合聯合變異量與速度輸送代理量的主導條件。",
            "海面高度呈正異常，東半部以南向流速異常最明顯，西南側局地可弱化或反向；PC 月平均由冬季負相位轉為暖季正相位，顯示可能存在季節性輸送調制。",
            "主要為東向至東北向異常，西側帶有旋轉或轉向結構，海面高度正異常由西北向東南減弱；可視為近岸—外海差異反應的次要模式。",
            "北部與東北部局地流向及海面高度梯度較突出，空間尺度縮小且速度同向性下降；其物理解讀應保留為區域剪切或邊界調整的高階候選。",
            "海面高度呈西北負、東南正的偶極，水平流向亦存在相反區塊；此類高階偶極可能對局地滯留或分流有影響，但整體變異貢獻有限。",
        ),
    ),
    RegionSpec(
        key="guishan",
        name_zh="龜山島",
        run_id="guishan_surface_u_v_eta_available_2024_2025_v1_80277f475e1f",
        location_note="龜山島及臺灣東北外海",
        status_zh="候選試驗區（candidate_pilot）",
        transport_choice="統計上 Mode 1 為聯合狀態主模態；若目標限定為水平海廢輸送，Mode 2 的速度 RMS 更大且同向性更高，列為本區速度主導候選。",
        overall_interpretation="本區變異量分散於前三模態，顯示流場結構較不易由單一方向概括。Mode 1 受海面高度成分影響較大，Mode 2 與 Mode 3 對水平速度輸送的貢獻更具直接性。",
        mode_texts=(
            "海面高度為廣域負異常，水平速度僅呈較弱的北至東北向訊號，且東半部較強；它是聯合狀態的第一模態，但不是前五模態中速度異常最強者。",
            "流場為高度一致的東北向異常，海面高度為正，速度 RMS 與空間同向性均優於 Mode 1；就水平海廢輸送而言，此模態為本區主要候選。",
            "主要呈東至東南向異常，海面高度整體較弱且島嶼附近帶有小尺度正負差異；其解釋變異量仍達雙位數，顯示為不可忽略的次要速度模式。",
            "南北區塊出現相反或轉向的流速結構，海面高度訊號小，速度同向性偏低；可能代表旋轉、剪切或傳播訊號的一組高階分量。",
            "呈多區塊反向流與局地化海面高度結構，適合視為細尺度殘差模態；應與 Mode 4 共同檢查是否形成成對傳播型態，而非單獨作定常機制解讀。",
        ),
    ),
    RegionSpec(
        key="houwan_nmmba",
        name_zh="後灣（海生館）",
        run_id="houwan_nmmba_surface_u_v_eta_available_2024_2025_v1_4ed000d7a3a4",
        location_note="恆春半島西側後灣及海生館鄰近海域",
        status_zh="候選試驗區（candidate_pilot）",
        transport_choice="背景平均流可能先決定淨輸送；Mode 1 為主要變率／反轉候選，Mode 2 則與平均東南向流同向且季節性明顯，是強化或弱化淨輸送的關鍵調制候選。",
        overall_interpretation="本區具有約 0.371 m s⁻¹ 的強勁東南向面積平均流，故候選模態必須在背景流脈絡下解讀。正相位的 Mode 1 傾向抵銷或反轉平均東南向流；Mode 2 正相位則強化背景輸送。",
        mode_texts=(
            "正相位為北至西北向異常並伴隨正海面高度，與強勁的東南向平均流相反；其重要性在於改變、減弱甚至局部反轉背景輸送，而非單純代表最常見的絕對流向。",
            "正相位為一致的東南向異常並伴隨正海面高度，與平均流同向；PC 具有鮮明季節相位，因此可能控制背景東南向輸送在不同季節的增強與減弱。",
            "以東至東北向異常為主，東西側存在轉向，海面高度多為負；它可造成沿岸與灣外之間的方向偏轉，屬第三順位但仍具雙位數變異貢獻的模式。",
            "回歸場集中於局部剪切與反向區塊，空間整合後的輸送訊號較弱；其代表的是灣口附近的局地交換或滯留型態。",
            "屬低能量、高空間梯度的殘餘結構，對整體淨輸送的貢獻有限；若與海岸幾何或潮相一致，仍可能影響短時滯留，但本分析不足以單獨證實。",
        ),
    ),
    RegionSpec(
        key="hsinchu",
        name_zh="新竹",
        run_id="hsinchu_surface_u_v_eta_available_2024_2025_v1_99801451634b",
        location_note="新竹沿岸及臺灣海峽東側",
        status_zh="候選試驗區（candidate_pilot）",
        transport_choice="Mode 1 為明確主要候選；其東北向速度異常與平均流同向、同向性極高。Mode 2 雖解釋變異量高，但主要由海面高度成分支配。",
        overall_interpretation="前兩模態已解釋絕大多數聯合變異；然而第二模態的速度幅度很小，顯示聯合變異量排名不能直接等同於漂浮海廢的水平輸送排名。",
        mode_texts=(
            "正相位為幾乎全區一致的東北向異常並伴隨負海面高度，且與東北向平均流同向；其速度同向性接近一，為本區最具一致性的水平輸送候選。",
            "海面高度為幅度很大的廣域負異常，水平速度僅在局部呈微弱南向或西南向；雖然聯合解釋變異量居第二，其直接水平輸送能力遠低於 Mode 1。",
            "西側偏東北、東側偏南的反向結構顯示近岸—外海剪切，海面高度為較弱正異常；可能影響分流，但整體變異量已小。",
            "速度與海面高度均為局地化高階結構，對全域水平輸送的整合訊號有限；物理解釋需倚賴獨立資料或重抽樣穩健性。",
            "為前五模態中最弱的局部殘差結構，可能受邊界、網格解析度與短尺度波動影響，不宜作為主要海廢傳輸機制。",
        ),
    ),
    RegionSpec(
        key="nangan",
        name_zh="南竿",
        run_id="nangan_surface_u_v_eta_available_2024_2025_v1_8661ce6730ad",
        location_note="馬祖南側島嶼群及鄰近水道",
        status_zh="已核定分析區（analysis_ready）",
        transport_choice="Mode 1 為主要候選；其一致東向速度異常最強。Mode 2 為東至東南向次要模式，Mode 3 則代表北向擾動。",
        overall_interpretation="前三模態分別概括本區東向、東南向與北向速度異常；Mode 1 為最強且最一致的水平輸送候選，Mode 2 與 Mode 3 描述次要方向調制。",
        mode_texts=(
            "正相位為全區一致的東向異常並伴隨負海面高度，速度幅度與同向性均為本區最高；此模式最可能顯著改變島間水道與區域東西向輸送。",
            "以東至東南向異常為主且海面高度為正，流速較 Mode 1 弱但仍高度同向；可視為第二順位的廣域輸送模式。",
            "呈一致北向異常並伴隨弱正海面高度，是前兩個東向模式以外的重要正交方向；可能在特定相位造成南北向偏轉。",
            "主要為島嶼周邊的局部反向與補償結構，速度與變異貢獻均低；其海廢意義較可能侷限於近岸滯留或短距離交換。",
            "為低能量的細尺度反向結構，對區域整體輸送排序不具主導性；需避免對高階空間紋理作過度物理解釋。",
        ),
    ),
)

# 正式報告沿用六區 batch 設定與計畫書的呈現順序，確保表格、章節、圖號及 provenance
# 可逐一對照；RegionSpec 的宣告位置僅為程式組織，不參與科學排序。
REPORT_REGION_ORDER = ("guishan", "gongliao", "hsinchu", "houwan_nmmba", "beigan", "nangan")


def parse_args() -> argparse.Namespace:
    """解析輸出位置；所有科學輸入固定由版本化成果目錄取得。"""

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=OUTPUT_DEFAULT, help="輸出 DOCX 路徑")
    return parser.parse_args()


def sha256_file(path: Path) -> str:
    """以串流方式計算檔案 SHA-256，避免將大型 PNG 或陣列整份載入記憶體。"""

    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_json(path: Path) -> dict[str, object]:
    """讀取 UTF-8 JSON；錯誤直接向上傳遞，避免缺少 provenance 時仍產生報告。"""

    return json.loads(path.read_text(encoding="utf-8"))


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, *, math_font: bool = False) -> None:
    """直接設定 run 的四組 Word 字型欄位，避免標題或英文被主題字型覆寫。

    Word 會分別保存 ASCII、High ANSI、東亞文字與複雜文字的字型。只寫 ``font.name`` 或
    ``w:eastAsia`` 其中一項時，標題中的英文字母或部分中文仍可能套用佈景主題字型；因此
    四項一律明確指定為「標楷體-繁」。``math_font`` 參數保留呼叫介面相容性，但依本版
    格式要求，數學段落也使用同一字族。
    """

    font_name = MATH_FONT if math_font else LATIN_FONT
    run.font.name = font_name
    r_fonts = run._element.get_or_add_rPr().rFonts
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    """設定表格儲存格底色；僅用於欄首與關鍵結論列，避免過度裝飾。"""

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    """套用 skill 指定的固定儲存格內距，單位為 twentieths of a point。"""

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    """固定欄寬，避免 Word 自動配寬造成欄位在不同引擎間漂移。"""

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_layout(table, widths: Sequence[int]) -> None:
    """依總寬 9360 DXA 建立固定表格幾何，並套用一致內距與垂直置中。"""

    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"表格欄寬總和必須為 {TABLE_WIDTH_DXA} DXA，實際為 {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    # ``tcW`` 只約束個別儲存格，LibreOffice 仍可能依單列文字長度重算視覺欄線；同步
    # 更新 ``tblGrid/gridCol`` 才能讓所有列共享同一欄界，尤其避免兩位數章號的目錄列
    # 把頁碼欄拉回文字尾端。
    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    if len(grid_columns) != len(widths):
        raise ValueError(f"表格 grid 欄數 {len(grid_columns)} 與指定欄寬數 {len(widths)} 不一致")
    for grid_column, width in zip(grid_columns, widths, strict=True):
        grid_column.set(qn("w:w"), str(width))
    for row in table.rows:
        # 表格列不得在兩頁間拆成上下兩半；對單列 code block 而言，這也能保證整組重現
        # 命令移到下一頁完整呈現，不會只把尾端參數孤立到後頁。
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table_text(table, header: bool = True, font_size: float = 9.0) -> None:
    """統一表格字級與欄首樣式；內容保留完整數字，不以縮字隱藏資訊。"""

    if header and table.rows:
        # 跨頁表格在每頁重複欄首，讓尺度常數或 provenance 的續頁仍可獨立判讀欄位。
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if header and row_index == 0:
                set_cell_shading(cell, NAVY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, font_size, bold=(header and row_index == 0), color=WHITE if header and row_index == 0 else None)


def add_table_caption(document: Document, text: str) -> None:
    """新增學術表說；表號由呼叫端明確給定，以確保跨頁與轉檔後仍穩定。"""

    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, 9.0, bold=False, color=MID_GRAY)


def add_figure(document: Document, path: Path, caption: str, width_inches: float) -> None:
    """以 inline picture 插入 PNG 並加圖說；若缺檔即停止，禁止靜默略過比例尺圖。"""

    if not path.is_file():
        raise FileNotFoundError(f"缺少報告圖檔：{path}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = document.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, 9.0, color=MID_GRAY)


def add_field(run, instruction: str) -> None:
    """在 run 內加入 Word 欄位，例如 PAGE、NUMPAGES 或 TOC。"""

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, end))


def add_code_block(document: Document, lines: Iterable[str]) -> None:
    """以單欄灰底表格呈現可複製命令；內容不換成圖片，便於讀者重現。"""

    table = document.add_table(rows=1, cols=1)
    set_table_layout(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GRAY)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Menlo"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
        run.font.size = Pt(8.0)


def configure_document(document: Document) -> None:
    """建立 Letter portrait 學術報告版面、字級階層與頁首頁尾。"""

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal._element.rPr.rFonts.set(qn(f"w:{attribute}"), CJK_FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Title": (24, NAVY, 12, 12),
        "Subtitle": (12, MID_GRAY, 4, 8),
        "Heading 1": (16, BLUE, 12, 8),
        "Heading 2": (13, NAVY, 10, 5),
        "Heading 3": (11, BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = LATIN_FONT
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{attribute}"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = LATIN_FONT
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        caption._element.rPr.rFonts.set(qn(f"w:{attribute}"), CJK_FONT)
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = False

    if "Equation" not in styles:
        equation = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = styles["Equation"]
    equation.font.name = MATH_FONT
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        equation._element.rPr.rFonts.set(qn(f"w:{attribute}"), MATH_FONT)
    equation.font.size = Pt(10.5)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(4)

    # 正式交付版不設頁首；保留既定上邊界，避免移除文字後正文與已校對頁碼整體上移。
    # section.header 仍由 DOCX 保留空白容器，以確保 Word 與 LibreOffice 的分頁結果一致。
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("第 ")
    set_run_font(footer_run, 8.5, color=MID_GRAY)
    add_field(footer_run, "PAGE")
    suffix = footer_p.add_run(" 頁，共 ")
    set_run_font(suffix, 8.5, color=MID_GRAY)
    add_field(suffix, "NUMPAGES")
    end = footer_p.add_run(" 頁")
    set_run_font(end, 8.5, color=MID_GRAY)


def enforce_document_font(document: Document) -> None:
    """在儲存前對所有既有文字 run 套用「標楷體-繁」直接格式。

    樣式層的字型設定仍可能被 Word 佈景主題或語系規則替代，尤其是 Title、Heading 與表格
    內新建立的 run。此函式遍歷正文、巢狀表格、頁首及頁尾，僅補上字型名稱，不改變原有
    字級、粗體、顏色或段落配置；因此正文的 12 pt 黑色與標題階層仍各自保留。
    """

    def apply_to_table(table) -> None:
        """遞迴處理表格儲存格及可能存在的巢狀表格。"""

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run)
                for nested_table in cell.tables:
                    apply_to_table(nested_table)

    for paragraph in document.paragraphs:
        is_body_text = paragraph.style.name in {"Normal", "List Bullet", "List Number"}
        for run in paragraph.runs:
            set_run_font(run)
            # 正文與條列內容一律採 12 pt 黑色；標題、公式、圖說、頁首頁尾及表格各自保留
            # 既定階層字級，以免 12 pt 強制套用後破壞章節導覽或表格版面。
            if is_body_text and run.text:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor.from_string("000000")
    for table in document.tables:
        apply_to_table(table)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)
            for table in part.tables:
                apply_to_table(table)


def add_cover(document: Document) -> None:
    """建立具留白、研究工項與版本資訊的 editorial cover。"""

    for _ in range(5):
        document.add_paragraph()
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("指定海域表層流場模態萃取與時序列變化\n主導海廢輸送候選模態報告")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("工項：流場模態萃取與時序列變化")
    document.add_paragraph()
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    run = rule.add_run("━" * 20)
    set_run_font(run, 9, color=BLUE)
    details = (
        "分析變數｜表層東向流速 u、北向流速 v、海表面高度 η\n"
        "分析期間｜2024–2025 年全部可得逐時樣本\n"
        "分析區域｜龜山島、貢寮、新竹、後灣（海生館）、北竿、南竿\n"
        "圖件版本｜academic_report_ready_v8（空間圖均含向量比例尺）\n"
        "報告版本｜聚焦版 2.0｜2026 年 8 月 4 日"
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.add_run(details)
    # 僅保留兩行留白，確保強制分頁符仍落在封面頁內；留白過多會讓 LibreOffice 把分頁符
    # 推到下一頁，進而在目錄前產生無內容的空白頁。
    for _ in range(2):
        document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_run = note.add_run(
        "研究定位：以多變量奇異值分解萃取各研究海域主要空間模態與標準化時間係數，"
        "並在各區自身流場脈絡下判讀主導海廢輸送候選模態。"
    )
    set_run_font(note_run, 9.5, color=MID_GRAY)
    document.add_page_break()


def add_toc(document: Document) -> None:
    """加入已依最終渲染校對的靜態章節目錄。

    headless LibreOffice 不會自動更新 Word TOC 欄位，會產生一頁只有標題的空白目錄；
    因此本報告使用可跨 Word／LibreOffice 穩定顯示的兩欄目錄。頁碼以 Letter 版面、
    標楷體-繁與目前 72 張圖件的最終渲染結果核定，若日後增刪內容應重新渲染並更新。
    """

    document.add_heading("目錄", level=1)
    entries = [
        ("摘要", "3"),
        ("1　研究背景、工項定位與研究問題", "3"),
        ("2　資料來源、研究區與品質控制", "3"),
        ("3　多變量 SVD 方法、採用理由與實作說明", "4"),
        ("4　六區 SVD 整體比較與主導候選判定", "7"),
        ("5　龜山島表層流場 SVD 成果與解讀", "9"),
        ("6　貢寮表層流場 SVD 成果與解讀", "17"),
        ("7　新竹表層流場 SVD 成果與解讀", "25"),
        ("8　後灣（海生館）表層流場 SVD 成果與解讀", "32"),
        ("9　北竿表層流場 SVD 成果與解讀", "40"),
        ("10　南竿表層流場 SVD 成果與解讀", "50"),
        ("11　跨區綜合、海廢輸送意涵與後續驗證", "57"),
        ("12　限制與不確定性", "57"),
        ("13　結論與建議", "58"),
        ("參考文獻", "58"),
        ("附錄 A　成果可追溯性與重現核對表", "59"),
    ]
    table = document.add_table(rows=0, cols=2)
    for label, page in entries:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = page
        cells[0].paragraphs[0].paragraph_format.space_after = Pt(2)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[1].paragraphs[0].paragraph_format.space_after = Pt(2)
    set_table_layout(table, [8200, 1160])
    format_table_text(table, header=False, font_size=9.5)
    document.add_page_break()


def add_focused_toc(document: Document) -> None:
    """加入聚焦版靜態目錄；頁碼須在最終逐頁渲染後回填。

    LibreOffice headless 不會可靠更新 Word 自動目錄，因此採固定兩欄表格。聚焦版只列出
    共用方法、六個彼此獨立的研究海域章節、各區結論與參考文獻，不放入跨區比較或附錄。
    """

    document.add_heading("目錄", level=1)
    entries = [
        ("摘要", "3"),
        ("1　研究背景、工項定位與研究問題", "3"),
        ("2　資料來源與品質控制", "3"),
        ("3　多變量 SVD 方法、採用理由與實作說明", "4"),
        ("4　龜山島表層流場 SVD 成果與解讀", "8"),
        ("5　貢寮表層流場 SVD 成果與解讀", "16"),
        ("6　新竹表層流場 SVD 成果與解讀", "24"),
        ("7　後灣（海生館）表層流場 SVD 成果與解讀", "32"),
        ("8　北竿表層流場 SVD 成果與解讀", "40"),
        ("9　南竿表層流場 SVD 成果與解讀", "49"),
        ("10　各研究海域之候選模態結論與方法限制", "56"),
        ("參考文獻", "58"),
    ]
    table = document.add_table(rows=0, cols=2)
    for label, page in entries:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = page
        cells[0].paragraphs[0].paragraph_format.space_after = Pt(2)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[1].paragraphs[0].paragraph_format.space_after = Pt(2)
    set_table_layout(table, [8200, 1160])
    format_table_text(table, header=False, font_size=9.5)
    document.add_page_break()


def add_equation(document: Document, equation: str, number: int) -> None:
    """以可搜尋文字呈現公式，保留編號並避免轉成不可編輯圖片。"""

    paragraph = document.add_paragraph(style="Equation")
    run = paragraph.add_run(f"{equation}    ({number})")
    set_run_font(run, 10.5, math_font=True)


def load_region_data(spec: RegionSpec) -> dict[str, object]:
    """載入單區小型 SVD 成果並計算報告用統計。

    面積平均僅使用 ``valid_mask`` 內的海洋格點，權重為 ``cell_area_m2``。速度模態強度
    定義為每 1 個標準差 PC 下 ``sqrt(<u_k²+v_k²>_area)``；同向性定義為面積平均向量
    大小除以該 RMS，介於 0 與 1，僅衡量方向一致性，不代表實際海廢通量。
    """

    run_dir = RESULT_ROOT / "svd" / spec.run_id
    bundle_dir = RESULT_ROOT / "svd_figure_bundles" / spec.run_id / "academic_report_ready_v8"
    metadata = load_json(run_dir / "metadata.json")
    bundle_metadata = load_json(bundle_dir / "metadata.json")
    mask = np.load(run_dir / "valid_mask.npy", allow_pickle=False).astype(bool)
    area = np.load(run_dir / "cell_area_m2.npy", allow_pickle=False)
    weights = area[mask]
    weights = weights / np.sum(weights)

    def area_mean(array: np.ndarray) -> float:
        return float(np.sum(array[mask] * weights))

    mean_u = np.load(run_dir / "mean_u.npy", allow_pickle=False)
    mean_v = np.load(run_dir / "mean_v.npy", allow_pickle=False)
    mean_eta = np.load(run_dir / "mean_eta.npy", allow_pickle=False)
    reg_u = np.load(run_dir / "regression_u.npy", allow_pickle=False)[:5]
    reg_v = np.load(run_dir / "regression_v.npy", allow_pickle=False)[:5]
    reg_eta = np.load(run_dir / "regression_eta.npy", allow_pickle=False)[:5]
    explained = np.load(run_dir / "explained_variance.npy", allow_pickle=False)[:5]
    cumulative = np.load(run_dir / "cumulative_explained_variance.npy", allow_pickle=False)[:5]
    pc = np.load(run_dir / "pc_standardized.npy", mmap_mode="r", allow_pickle=False)[:5]
    time_ns = np.load(run_dir / "time_utc_ns.npy", mmap_mode="r", allow_pickle=False)
    imputed = np.load(run_dir / "imputed_mask.npy", mmap_mode="r", allow_pickle=False)

    mode_metrics: list[dict[str, float]] = []
    for index in range(5):
        speed = np.sqrt(reg_u[index][mask] ** 2 + reg_v[index][mask] ** 2)
        velocity_rms = math.sqrt(float(np.sum((speed ** 2) * weights)))
        mean_vector = math.hypot(area_mean(reg_u[index]), area_mean(reg_v[index]))
        mode_metrics.append(
            {
                "mean_u": area_mean(reg_u[index]),
                "mean_v": area_mean(reg_v[index]),
                "mean_eta": area_mean(reg_eta[index]),
                "velocity_rms": velocity_rms,
                "velocity_q95": float(np.quantile(speed, 0.95)),
                "directional_coherence": mean_vector / velocity_rms if velocity_rms else 0.0,
                "pc_mean": float(np.mean(pc[index])),
                "pc_std": float(np.std(pc[index], ddof=1)),
            }
        )

    # time_utc_ns 與 PC shape 是重現時序列的關鍵資料契約；此處主動檢查而非假設一致。
    if pc.shape[1] != time_ns.shape[0]:
        raise ValueError(f"{spec.name_zh} PC 與時間軸長度不一致")
    if int(np.count_nonzero(imputed)) != 0:
        raise ValueError(f"{spec.name_zh} 與既有 QC 摘要不符：偵測到插補值")

    return {
        "spec": spec,
        "run_dir": run_dir,
        "bundle_dir": bundle_dir,
        "metadata": metadata,
        "bundle_metadata": bundle_metadata,
        "mask": mask,
        "area_km2": float(np.sum(area[mask]) / 1_000_000.0),
        "cell_count": int(np.count_nonzero(mask)),
        "sample_count": int(time_ns.shape[0]),
        "mean_u": area_mean(mean_u),
        "mean_v": area_mean(mean_v),
        "mean_eta": area_mean(mean_eta),
        "mean_speed": math.hypot(area_mean(mean_u), area_mean(mean_v)),
        "explained": explained,
        "cumulative": cumulative,
        "mode_metrics": mode_metrics,
        "velocity_scale": float(metadata["svd"]["velocity_rms_mps"]),
        "eta_scale": float(metadata["svd"]["eta_rms_m"]),
        "imputed_count": 0,
    }


def add_summary_chapters(document: Document, data: Sequence[dict[str, object]], counters: dict[str, int]) -> None:
    """撰寫摘要、研究目的、資料與跨區比較，先交代最重要結論再進入技術細節。"""

    document.add_heading("摘要", level=1)
    document.add_paragraph(
        "本報告對龜山島、貢寮、新竹、後灣（海生館）、北竿與南竿六指定海域之 2024–2025 年表層流場，"
        "以東向流速 u、北向流速 v 與海表面高度 η 建立面積加權、多變量 SVD。研究目的為萃取主要空間特徵及其"
        "時序列，並依聯合解釋變異量、水平速度異常強度、空間同向性、背景平均流與季節持續性，提出主導海廢輸送的"
        "優先驗證候選模態。六區資料可得率介於 97.61%–97.74%，共同有效海洋格點保留後均無短缺口插補；前五模態"
        "累積解釋 86.07%–98.84% 的聯合變異。"
    )
    document.add_paragraph(
        "綜合結果顯示，北竿、南竿、貢寮與新竹均以 Mode 1 為最主要水平輸送候選；後灣 Mode 1 控制對強東南向"
        "平均流的減弱或反轉，而 Mode 2 控制同向增強，故淨輸送仍可能由背景平均流優先支配；龜山島的 Mode 1 雖"
        "為聯合變異量最大模態，速度能量卻低於 Mode 2，因此若判準限定為漂浮海廢的水平平流，Mode 2 應列為優先"
        "驗證候選。上述判定屬資料驅動的模態篩選，不等同於海廢軌跡或來源歸因；最終結論需以重建流場的粒子追蹤、"
        "截面通量與外力敏感度試驗完成驗證。"
    )
    keywords = document.add_paragraph()
    key_run = keywords.add_run("關鍵詞：")
    set_run_font(key_run, bold=True)
    keywords.add_run("海洋流場、奇異值分解、經驗正交函數、多變量模態、表層海流、海洋廢棄物、拉格朗日輸送")

    document.add_heading("1　研究背景、工項定位與研究問題", level=1)
    document.add_paragraph(
        "本工項「流場模態萃取與時序列變化」針對指定熱點海域之表層海流資料，利用奇異值分解進行空間特徵萃取，"
        "並進一步探討主要流場模態的時間變化。其最終應用目標不是僅列示統計變異排序，而是判定哪些模態最可能"
        "控制漂浮海洋廢棄物的輸送方向、強度、跨區交換或近岸滯留。由於海廢平流直接受水平速度控制，而聯合 SVD"
        "同時納入 u、v 與 η，本報告明確區分「聯合狀態變異量最大」與「水平速度輸送最重要」兩種判準。"
    )
    questions = [
        "六區表層 u–v–η 聯合變異可由多少個主要模態概括？",
        "各模態的空間流向、海面高度關係與 PC 時序特徵為何？",
        "在納入速度異常強度、方向一致性及平均流後，哪一模態最值得優先作為海廢輸送驗證對象？",
        "資料缺口、尺度正規化、區域範圍與模態正交性對結論有何限制？",
    ]
    for item in questions:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("2　資料來源、研究區與品質控制", level=1)
    document.add_heading("2.1　資料來源與分析變數", level=2)
    document.add_paragraph(
        "科學輸入為 OCM 表層快取之逐時 u（m s⁻¹）、v（m s⁻¹）與 η（m），期間涵蓋 2024 年 1 月至 2025 年 12 月"
        "的全部可得樣本。各區依版本化局部研究範圍的格點中心納入規則裁切，先套用靜態海洋遮罩，再要求同一格點"
        "的 u、v、η 三者在至少 95% 時間共同有效。短缺口僅允許在前後均有有效觀測且長度不超過 2 小時時作線性插值；"
        "插值後任何仍不完整的時間步均整筆移除，從未以零值填補。六區實際插補筆數均為零。"
    )

    headers = ["區域", "分析範圍（lon, lat）", "有效格點", "面積（km²）", "樣本數", "可得率", "狀態"]
    table = document.add_table(rows=1, cols=len(headers))
    table.rows[0].cells[0].text = headers[0]
    for index, header in enumerate(headers[1:], start=1):
        table.rows[0].cells[index].text = header
    for region in data:
        metadata = region["metadata"]
        bbox = metadata["analysis_unit"]["analysis_bbox_lon_lat"]
        row = table.add_row().cells
        row[0].text = region["spec"].name_zh
        row[1].text = f"{bbox[0]:.2f}–{bbox[1]:.2f}; {bbox[2]:.2f}–{bbox[3]:.2f}"
        row[2].text = f"{region['cell_count']:,}"
        row[3].text = f"{region['area_km2']:.3f}"
        row[4].text = f"{region['sample_count']:,}"
        coverage = region["sample_count"] / 17544 * 100
        row[5].text = f"{coverage:.2f}%"
        row[6].text = region["spec"].status_zh
    set_table_layout(table, [900, 2100, 900, 1050, 900, 900, 2610])
    format_table_text(table, font_size=7.7)
    counters["table"] += 1
    add_table_caption(document, f"表 {counters['table']}　六分析區之空間範圍、有效格點與可得樣本。分母 17,544 為 2024–2025 兩年逐時理論樣本數。")

    document.add_heading("2.2　時間軸標準化與缺口品質", level=2)
    document.add_paragraph(
        "月份檔案合併後先依 UTC 排序，重複時間戳採 prefer-last 去重。貢寮與龜山島另依版本化 known-time-axis repair"
        "將 2025 年 7 月首 24 筆錯置時間平移 24 小時；修正僅改變時間座標，不改寫 u、v 或 η 值。貢寮與龜山島最終"
        "各保留 17,148 筆，47 個時間斷點；其餘四區各保留 17,124 筆，48 個斷點。六區最大資料間隙均為 50 小時，"
        "因此 PC 圖中的空白或不連續時段應視為真實資料缺口，而非被插值平滑的訊號。"
    )

    document.add_heading("2.3　平均流場", level=2)
    headers = ["區域", "面積平均 u", "面積平均 v", "平均向量速率", "平均 η", "流向概述"]
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for region in data:
        u = region["mean_u"]
        v = region["mean_v"]
        direction = vector_direction_zh(u, v)
        row = table.add_row().cells
        row[0].text = region["spec"].name_zh
        row[1].text = f"{u:+.4f} m s⁻¹"
        row[2].text = f"{v:+.4f} m s⁻¹"
        row[3].text = f"{region['mean_speed']:.4f} m s⁻¹"
        row[4].text = f"{region['mean_eta']:.4f} m"
        row[5].text = direction
    set_table_layout(table, [1050, 1500, 1500, 1650, 1450, 2210])
    format_table_text(table, font_size=8.2)
    counters["table"] += 1
    add_table_caption(document, f"表 {counters['table']}　六區面積加權平均表層流場。平均向量速率為 |〈u〉,〈v〉|，不等同於格點流速大小的面積平均。")


def add_focused_summary_chapters(document: Document) -> None:
    """撰寫聚焦版摘要、工項定位與共用資料處理契約。

    聚焦版只說明六個獨立研究海域共用的科學方法，不建立合併統計表或區域間排序；
    各區的有效格點、正規化尺度、平均場及候選模態均在對應章節獨立呈現。
    """

    document.add_heading("摘要", level=1)
    document.add_paragraph(
        "本報告針對龜山島、貢寮、新竹、後灣（海生館）、北竿與南竿之 2024–2025 年表層流場，"
        "分別以東向流速 u、北向流速 v 與海表面高度 η 建立面積加權、多變量 SVD。工項重點為"
        "「流場模態萃取與時序列變化」：對每一研究海域獨立萃取主要空間模態，解析標準化主成分（PC）"
        "的逐時變化、30 日平滑與月平均相位，並依該區的聯合解釋變異量、水平速度異常 RMS、方向同向性"
        "及背景平均流，界定主導海廢輸送候選模態。"
    )
    document.add_paragraph(
        "報告完整呈現每一海域之平均場、解釋變異量譜及前五模態的物理回歸場與 PC 時序，並於各自"
        "章節說明正、負相位所代表的流向變化、季節調制及候選角色。所有判讀均限定於該研究海域自身的"
        "空間結構、背景平均流與時間變化脈絡，不建立區域間排序或合併推論。"
    )
    keywords = document.add_paragraph()
    key_run = keywords.add_run("關鍵詞：")
    set_run_font(key_run, bold=True)
    keywords.add_run("海洋流場、奇異值分解、多變量模態、主成分時序、表層海流、海洋廢棄物")

    document.add_heading("1　研究背景、工項定位與研究問題", level=1)
    document.add_paragraph(
        "本工項「流場模態萃取與時序列變化」針對指定熱點海域之表層海流資料，利用奇異值分解進行空間"
        "特徵萃取，並探討主要流場模態的時間變化。研究輸出不僅列示統計變異排序，也判讀哪些模態最可能"
        "控制漂浮海洋廢棄物在各自研究海域內的輸送方向、強度與近岸滯留。由於海廢平流直接受水平速度"
        "控制，而聯合 SVD 同時納入 u、v 與 η，本報告明確區分「聯合狀態變異量最大」與「水平速度輸送"
        "最重要」兩種判準。"
    )
    questions = [
        "各研究海域的表層 u–v–η 聯合變異可由多少個主要模態概括？",
        "各海域前五模態的空間流向、海面高度關係與 PC 時序特徵為何？",
        "在納入該區速度異常強度、方向一致性及平均流後，哪一模態為主導海廢輸送候選？",
        "資料缺口、尺度正規化、區域範圍與模態正交性如何限制各海域的獨立判讀？",
    ]
    for item in questions:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("2　資料來源與品質控制", level=1)
    document.add_heading("2.1　資料來源與分析變數", level=2)
    document.add_paragraph(
        "科學輸入為 OCM 表層快取之逐時 u（m s⁻¹）、v（m s⁻¹）與 η（m），期間涵蓋 2024 年 1 月至"
        " 2025 年 12 月的全部可得樣本。每一研究海域依版本化局部研究範圍，以格點中心是否落在"
        "範圍內作為裁切規則，"
        "先套用靜態海洋遮罩，再要求同一格點的 u、v、η 三者在至少 95% 時間共同有效。短缺口僅允許在"
        "前後均有有效觀測且長度不超過 2 小時時作線性插值；插值後任何仍不完整的時間步均整筆移除，"
        "從未以零值填補。本次各海域實際插補筆數均為零。"
    )
    document.add_heading("2.2　時間軸標準化與缺口品質", level=2)
    document.add_paragraph(
        "月份檔案合併後先依 UTC 排序，重複時間戳採 prefer-last 去重。貢寮與龜山島另依版本化"
        " known-time-axis repair 將 2025 年 7 月首 24 筆錯置時間平移 24 小時；修正僅改變時間座標，"
        "不改寫 u、v 或 η 值。所有 PC 圖均保留資料缺口，不跨缺口插值；時序中的空白或不連續區段代表"
        "來源資料未提供樣本。"
    )
    document.add_heading("2.3　獨立研究海域之呈現原則", level=2)
    document.add_paragraph(
        "六個研究海域採相同資料契約與 SVD 實作，但局部研究範圍、有效格點、正規化尺度、平均流場、模態"
        "排序與 PC 均依該海域自身資料獨立計算。後續章節依序呈現單一海域的基本資料、平均場、解釋"
        "變異量與 Mode 1–5，不把不同海域的 Mode 編號、EV 或速度幅度建立共同排名。"
    )


def vector_direction_zh(u: float, v: float) -> str:
    """將東、北分量轉為八方位流去向；此處描述水體移動方向，不是來向。"""

    angle = (math.degrees(math.atan2(v, u)) + 360.0) % 360.0
    labels = ("東向", "東北向", "北向", "西北向", "西向", "西南向", "南向", "東南向")
    return labels[int((angle + 22.5) // 45.0) % 8]


def add_method_chapter(document: Document, counters: dict[str, int]) -> None:
    """以白話定義搭配公式說明 SVD 實作，使讀者能理解資料如何轉成模態與時序。"""

    document.add_heading("3　多變量 SVD 方法、採用理由與實作說明", level=1)
    document.add_heading("3.1　SVD 的作用與採用理由", level=2)
    document.add_paragraph(
        "本研究把每一個時間點的流場視為一筆資料。每筆資料同時包含所有海洋格點的東向流速 u、"
        "北向流速 v 與海表面高度 η，因此原始資料表具有大量彼此相關的欄位。奇異值分解（SVD）的"
        "作用，是從這張大型資料表中找出少數幾種反覆出現的共同變化型態。每一種型態稱為一個模態；"
        "模態包含一張空間分布圖，以及一條描述該空間型態何時增強、減弱或反向的時間係數。"
    )
    document.add_paragraph(
        "u、v 與 η 採聯合分析，而非分開計算，原因是同一次海況變化可能同時改變流向、流速與海面"
        "高度。聯合 SVD 可讓同一條時間係數同步連結三個變數，便於判讀速度異常與水位異常是否同時"
        "發生。不同模態在數學上彼此正交；白話而言，後一模態應描述前面模態尚未概括的變化，避免"
        "重複計算相同訊號。此方法不需預先指定季節函數或假設固定流向，適合用於探索主要流場變率。"
    )
    document.add_paragraph(
        "格點面積亦納入權重，避免面積較小的格點與面積較大的格點被視為具有完全相同的空間代表性；"
        "u 與 v 則共用同一速度尺度，以保留水平向量的方向與幾何關係（Bretherton et al., 1992；"
        "de Oliveira Júnior et al., 2022；Song et al., 2025）。SVD 描述的是線性共同變化，傳播或"
        "旋轉訊號有時會分布於相鄰的一對模態（Volkov et al., 2022），因此模態仍須結合空間圖與"
        "時間係數一起解讀。"
    )

    document.add_heading("3.2　資料前處理與共同有效樣本", level=2)
    steps = [
        "依版本化局部研究範圍，以格點中心是否位於封閉範圍內決定納入格點，再套用靜態海洋遮罩，排除陸地格點。",
        "建立 u、v、η 的共同有效遮罩。所謂有效，是三個變數皆為有限數值；格點必須在至少 95% 的時間同時有效才予以保留。",
        "僅允許補齊長度不超過 2 小時、且缺口前後皆有有效值的內部短缺口；本次所有研究海域的實際插補筆數均為 0。",
        "若某一時間點在保留格點中仍含缺值，則移除該完整時間點。此作法確保 SVD 輸入矩陣不含缺值，也不以 0 假裝成實際觀測。",
        "每一格點分別減去 2024–2025 年保留期間的時間平均，得到 u′、v′、η′。撇號代表相對於平均狀態的偏差，稱為距平。",
    ]
    for item in steps:
        document.add_paragraph(item, style="List Number")
    document.add_paragraph(
        "平均流場與距平模態的角色不同：平均流場描述分析期間的背景流向；SVD 則只分析相對背景值的"
        "增減與轉向。解讀海廢輸送候選模態時，必須把模態異常與背景平均流放在同一脈絡中。"
    )

    document.add_heading("3.3　RMS 尺度、狀態矩陣與面積權重", level=2)
    document.add_paragraph(
        "令 N 代表保留下來的時間點數，P 代表共同有效的海洋格點數，aₚ 代表第 p 個格點所涵蓋的"
        "實際面積。由於流速以 m s⁻¹ 表示，海表面高度以 m 表示，兩者單位與典型數值大小不同，不能"
        "直接放入同一矩陣比較。因此先各自除以一個代表典型變動幅度的尺度。"
    )
    document.add_paragraph(
        "本研究以均方根（root mean square，RMS）表示典型變動幅度。計算方式是先把每個距平值平方，"
        "依格點面積求平均，再取平方根。平方可避免正、負距平互相抵銷，因此 RMS 可理解為『通常會偏離"
        "平均值多大』。速度尺度 U₀ 同時納入 u′ 與 v′；海表面高度尺度 E₀ 則由 η′ 計算："
    )
    add_equation(document, "U₀ = { ΣₜΣₚ aₚ[u′ₚ(t)²+v′ₚ(t)²] / [2NΣₚaₚ] }¹ᐟ²", 1)
    add_equation(document, "E₀ = { ΣₜΣₚ aₚη′ₚ(t)² / [NΣₚaₚ] }¹ᐟ²", 2)
    document.add_paragraph(
        "除以 U₀ 或 E₀ 稱為尺度正規化。其目的不是改變物理資料，而是把速度與海面高度轉成可公平比較"
        "的無因次數值，避免其中一個變數只因單位或數值較大而支配結果。每一時間 t 的資料依固定順序"
        "排成狀態向量 x(t)：先放所有格點的 u′，再放 v′，最後放 η′。"
    )
    add_equation(document, "x(t) = [u′₁/U₀,…,u′ₚ/U₀, v′₁/U₀,…,v′ₚ/U₀, η′₁/E₀,…,η′ₚ/E₀]ᵀ", 3)
    document.add_paragraph(
        "把全部 N 個時間點並排後，即形成 3P × N 的資料矩陣。接著將每個格點乘以面積平方根 √aₚ。"
        "使用平方根的原因，是 SVD 後續會計算平方和；平方根經平方後恰好回到面積 aₚ，使每個格點對"
        "總變異的貢獻與其代表面積成正比。以 W 表示這組面積平方根權重，面積加權矩陣為："
    )
    add_equation(document, "Xw = W[x(t₁),…,x(tN)]", 4)

    document.add_heading("3.4　矩陣求解方式與解釋變異量", level=2)
    document.add_paragraph(
        "直接對 Xw 做 SVD 可以得到相同答案，但本資料的時間點數 N 約為 17,000，而空間變數數量 3P"
        "僅為 288–864。為節省記憶體與計算時間，程式先計算空間共變異矩陣 C。共變異矩陣可理解為一張"
        "『各位置與各變數是否經常一起增減』的關係表，其大小只與 3P 有關："
    )
    add_equation(document, "C = Xw Xwᵀ/(N−1) = UΛUᵀ", 5)
    document.add_paragraph(
        "對稱特徵分解把 C 拆成空間方向 U 與特徵值 Λ。U 的每一欄代表一個加權後的空間型態；特徵值"
        "λₖ 表示第 k 個型態所包含的變異量。程式將特徵值由大到小排序，所以 Mode 1 是聯合變異量最大"
        "的模態，Mode 2 是排除 Mode 1 後的第二大模態，依此類推。"
    )
    document.add_paragraph(
        "此作法與薄型 SVD 完全等價。『薄型』是指只保留資料真正需要的維度，不建立多餘的零空間。"
        "在 Xw=UΣVᵀ 中，U 表示空間型態，Σ 對角線上的奇異值 σₖ 表示模態強度，V 表示時間方向；"
        "主成分時間係數 PCₖ(t) 可由空間型態投影回每一時間點取得："
    )
    add_equation(document, "Xw = UΣVᵀ,   σₖ = √[λₖ(N−1)],   PCₖ(t) = UₖᵀXw", 6)
    document.add_paragraph(
        "解釋變異量（explained variance，EV）是單一模態變異量占全部模態變異量的比例。EV 越高，表示"
        "該模態越能概括 u、v、η 的整體共同變化；但若 EV 主要來自 η，並不必然代表水平流速也最強。"
        "本研究計算前 20 個模態，報告呈現前 5 個模態及其累積 EV。"
    )
    add_equation(document, "EVₖ = λₖ/Σⱼλⱼ = σₖ²/Σⱼσⱼ²", 7)

    document.add_heading("3.5　空間係數、物理回歸場與標準化 PC", level=2)
    document.add_paragraph(
        "SVD 最初得到的空間 loading，可譯為『空間係數』：它是一張係數地圖，說明各格點對某一模態"
        "的反應方向與相對大小；它本身不是某個特定時刻的實際流速。為恢復物理意義，程式先移除面積"
        "平方根權重，再乘回 U₀ 或 E₀，並乘上原始 PC 的標準差。如此得到的物理回歸場，就是報告空間圖"
        "中的箭頭與底色。"
    )
    document.add_paragraph(
        "PC（principal component，主成分時間係數）是一條無因次時間序列，表示該模態在每個時間點"
        "出現得多強。標準化後，PC 的平均值約為 0、標準差為 1；PC=+1 代表正相位達一個典型幅度，"
        "PC=−1 代表相同幅度的反相位。樣本標準差以 N−1 為分母，亦即程式中的 ddof=1。報告空間圖"
        "因此可直接解讀為 PC 改變 1 個標準差時，u、v 與 η 預期改變多少："
    )
    add_equation(document, "[u′ₖ,v′ₖ,η′ₖ](x,t) = Rₖ(x) · PC*ₖ(t)", 8)
    document.add_paragraph(
        "式中 Rₖ(x) 是第 k 模態的物理回歸場，PC*ₖ(t) 是標準化 PC。若 PC 為正，流速異常方向與"
        "圖中箭頭一致；若 PC 為負，箭頭方向與海面高度異常的正負必須同時反轉。SVD 的整體正負號"
        "在數學上可互換，因此程式使用固定符號規則，確保同一成果重算時不會任意翻轉。每一海域的"
        " Mode 編號只代表該海域內部的變異量排序。"
    )

    document.add_heading("3.6　主導海廢輸送候選模態的判讀方式", level=2)
    document.add_paragraph(
        "候選模態不只依 EV 排名。判讀時先看模態能概括多少聯合變異，再檢查速度異常本身是否夠強、"
        "箭頭是否大致同向、與背景平均流疊加後會增強或反轉何種流向，以及該 PC 相位是否具有持續性"
        "或季節規律。各項指標的白話意義如下。"
    )
    headers = ["判準", "計算或判讀方式", "白話意義", "限制"]
    rows = [
        ("聯合解釋變異量（EV）", "單一 EV 與累積 EV", "模態概括 u–v–η 整體變化的比例", "η 占比高時，不代表水平流速最強"),
        ("速度異常均方根（RMS）", "每 1 個標準差 PC 下的速度平方平均後開根號", "模態通常可造成多大的流速改變", "無法表達所有近岸邊界與局地渦旋"),
        ("方向同向性", "區域平均向量大小除以速度 RMS，介於 0 與 1", "接近 1 表示多數格點同向；接近 0 表示轉向或互相抵銷", "會受海岸形狀與來源位置影響"),
        ("背景平均流與模態疊加", "比較平均流與正、負一個標準差模態", "判斷模態會增強、削弱或反轉背景流", "極端 PC 不一定只停留在一個標準差"),
        ("PC 持續性與季節性", "檢查逐時 PC、30 日平滑與月平均", "判斷高影響相位何時出現、維持多久", "季節規律本身不等於驅動原因"),
    ]
    table = document.add_table(rows=1, cols=4)
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        row = table.add_row().cells
        for index, value in enumerate(values):
            row[index].text = value
    set_table_layout(table, [1800, 2600, 2700, 2260])
    format_table_text(table, font_size=8.5)
    counters["table"] += 1
    add_table_caption(document, f"表 {counters['table']}　由 SVD 模態判讀主導海廢輸送候選的分層判準。")


def add_comparison_chapter(document: Document, data: Sequence[dict[str, object]], counters: dict[str, int]) -> None:
    """建立跨區變異濃縮與海廢候選結論表。"""

    document.add_heading("4　六區 SVD 整體比較與主導候選判定", level=1)
    document.add_heading("4.1　解釋變異量濃縮程度", level=2)
    headers = ["區域", "Mode 1", "Mode 2", "Mode 3", "前 2 累積", "前 5 累積", "結構特徵"]
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for region in data:
        ev = region["explained"] * 100
        cum = region["cumulative"] * 100
        if cum[1] >= 85:
            feature = "高度集中於前兩模態"
        elif cum[1] >= 70:
            feature = "前兩模態為主、仍有次要模態"
        else:
            feature = "變異分散，前三模態均重要"
        row = table.add_row().cells
        values = [region["spec"].name_zh, f"{ev[0]:.2f}%", f"{ev[1]:.2f}%", f"{ev[2]:.2f}%", f"{cum[1]:.2f}%", f"{cum[4]:.2f}%", feature]
        for index, value in enumerate(values):
            row[index].text = value
    set_table_layout(table, [950, 900, 900, 900, 1050, 1050, 3610])
    format_table_text(table, font_size=8.2)
    counters["table"] += 1
    add_table_caption(document, f"表 {counters['table']}　六區前五模態解釋變異量比較。百分比以完整特徵值譜為分母。")
    document.add_paragraph(
        "新竹的前兩模態累積 95.18%，顯示聯合變異高度低維；北竿與南竿亦於前兩模態超過 86%。龜山島前兩模態僅"
        "61.10%，第三模態仍占 16.10%，代表其主要流場變率較分散。前五模態累積最低者仍為龜山島（86.07%），其餘"
        "各區均超過 90%。"
    )

    document.add_heading("4.2　水平速度代理量與候選模態", level=2)
    headers = ["區域", "M1 速度 RMS", "M2 速度 RMS", "M3 速度 RMS", "M1 同向性", "優先候選與判定"]
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for region in data:
        metrics = region["mode_metrics"]
        row = table.add_row().cells
        values = [
            region["spec"].name_zh,
            f"{metrics[0]['velocity_rms']:.3f}",
            f"{metrics[1]['velocity_rms']:.3f}",
            f"{metrics[2]['velocity_rms']:.3f}",
            f"{metrics[0]['directional_coherence']:.3f}",
            region["spec"].transport_choice,
        ]
        for index, value in enumerate(values):
            row[index].text = value
    set_table_layout(table, [900, 1200, 1200, 1200, 1100, 3760])
    format_table_text(table, font_size=7.8)
    counters["table"] += 1
    add_table_caption(document, f"表 {counters['table']}　每 1 個標準差 PC 的水平速度異常 RMS（m s⁻¹）及優先驗證候選。")
    document.add_paragraph(
        "跨區比較支持一項關鍵方法論結論：EV 衡量的是 u–v–η 聯合狀態，而非純速度輸送。新竹 Mode 2 的聯合 EV"
        "達 34.66%，速度 RMS 卻僅約 0.052 m s⁻¹；龜山島 Mode 2 的 EV 小於 Mode 1，但速度 RMS 約 0.171 m s⁻¹，"
        "高於 Mode 1 約 0.091 m s⁻¹。因此，主導海廢傳輸模式的篩選必須同時檢查物理速度回歸場。"
    )


def add_region_chapter(document: Document, region: dict[str, object], chapter_number: int, counters: dict[str, int]) -> None:
    """為單一研究海域建立可獨立閱讀的資料摘要、平均場、EV 與前五模態解讀。

    每章重述該區邊界、樣本覆蓋率、正規化尺度與平均流，使讀者不必引用其他海域章節
    即可理解該區 SVD 的輸入、尺度與候選結論；空間圖一律使用附向量比例尺的 v8 PNG。
    """

    spec: RegionSpec = region["spec"]
    bundle_dir: Path = region["bundle_dir"]
    fig_dir = bundle_dir / "figures/report"
    ev = region["explained"] * 100
    cum = region["cumulative"] * 100
    local_scope = region["metadata"]["analysis_unit"]["analysis_bbox_lon_lat"]
    coverage = region["sample_count"] / 17544 * 100
    heading = document.add_heading(f"{chapter_number}　{spec.name_zh}表層流場 SVD 成果與解讀", level=1)
    # 直接讓章標題在新頁開始，避免前一頁已接近滿版時，獨立分頁符被推到下一頁而形成空白頁。
    heading.paragraph_format.page_break_before = True
    document.add_paragraph(
        f"研究區位於{spec.location_note}，局部研究範圍為東經 {local_scope[0]:.2f}–{local_scope[1]:.2f}°、"
        f"北緯 {local_scope[2]:.2f}–{local_scope[3]:.2f}°，"
        f"區域狀態為「{spec.status_zh}」。共保留 {region['cell_count']:,} 個共同有效海洋格點，面積約 {region['area_km2']:.3f} km²；"
        f"逐時樣本 {region['sample_count']:,} 筆，可得率 {coverage:.2f}%。本區速度尺度 U₀={region['velocity_scale']:.6f} m s⁻¹，"
        f"海面高度尺度 E₀={region['eta_scale']:.6f} m。面積平均流為 u={region['mean_u']:+.4f}、v={region['mean_v']:+.4f} m s⁻¹，"
        f"合成向量速率 {region['mean_speed']:.4f} m s⁻¹，去向為{vector_direction_zh(region['mean_u'], region['mean_v'])}。"
        f"前兩模態累積解釋 {cum[1]:.2f}%，前五模態累積 {cum[4]:.2f}%。{spec.overall_interpretation}"
    )

    document.add_heading(f"{chapter_number}.1　平均場與變異量譜", level=2)
    counters["figure"] += 1
    add_figure(
        document,
        fig_dir / "mean_surface_flow_report_with_vector_scale.png",
        f"圖 {counters['figure']}　{spec.name_zh} 2024–2025 表層平均流場與平均海表面高度。箭頭比例尺位於圖內獨立圖例區；底色單位為 m。",
        5.05,
    )
    counters["figure"] += 1
    add_figure(
        document,
        fig_dir / "svd_explained_variance_report.png",
        f"圖 {counters['figure']}　{spec.name_zh} 表層 u–v–η 多變量 SVD 解釋變異量與累積比例。",
        6.1,
    )

    for mode_index, interpretation in enumerate(spec.mode_texts, start=1):
        metric = region["mode_metrics"][mode_index - 1]
        document.add_heading(f"{chapter_number}.{mode_index + 1}　Mode {mode_index}", level=2)
        role = "主要候選" if mode_index == 1 else "次要候選／高階結構"
        if spec.key == "guishan" and mode_index == 2:
            role = "水平速度主導候選"
        if spec.key == "houwan_nmmba" and mode_index == 2:
            role = "平均流同向調制候選"
        document.add_paragraph(
            f"Mode {mode_index} 解釋 {ev[mode_index - 1]:.2f}% 的聯合變異，累積達 {cum[mode_index - 1]:.2f}%。"
            f"每 1σ PC 的速度異常 RMS 為 {metric['velocity_rms']:.4f} m s⁻¹，格點速度第 95 百分位為 "
            f"{metric['velocity_q95']:.4f} m s⁻¹，方向同向性為 {metric['directional_coherence']:.3f}；"
            f"面積平均回歸分量為 u={metric['mean_u']:+.4f}、v={metric['mean_v']:+.4f} m s⁻¹、η={metric['mean_eta']:+.4f} m。"
            f"判定角色：{role}。{interpretation}"
        )
        counters["figure"] += 1
        add_figure(
            document,
            fig_dir / f"svd_mode_{mode_index:02d}_spatial_report_with_vector_scale.png",
            f"圖 {counters['figure']}　{spec.name_zh} Mode {mode_index} 之 1σ PC 物理回歸場。箭頭為流速異常（m s⁻¹），"
            "圖內含向量比例尺；底色為海表面高度異常（m）。PC 為負時，箭頭與底色符號同時反轉。",
            5.05,
        )
        counters["figure"] += 1
        add_figure(
            document,
            fig_dir / f"svd_mode_{mode_index:02d}_pc_report.png",
            f"圖 {counters['figure']}　{spec.name_zh} Mode {mode_index} 標準化 PC 時序、30 日平滑與月平均相位。空白時段保留資料缺口，不作跨缺口插值。",
            6.15,
        )


def add_independent_conclusions(document: Document, data: Sequence[dict[str, object]]) -> None:
    """逐區彙整候選模態與共用方法限制，不進行區域合併、排名或關聯推論。"""

    heading = document.add_heading("10　各研究海域之候選模態結論與方法限制", level=1)
    heading.paragraph_format.page_break_before = True
    for index, region in enumerate(data, start=1):
        spec: RegionSpec = region["spec"]
        ev = region["explained"] * 100
        cumulative = region["cumulative"] * 100
        document.add_heading(f"10.{index}　{spec.name_zh}", level=2)
        document.add_paragraph(
            f"{spec.transport_choice}本區 Mode 1 解釋 {ev[0]:.2f}% 的聯合變異，前兩模態累積為 "
            f"{cumulative[1]:.2f}%，前五模態累積為 {cumulative[4]:.2f}%。{spec.overall_interpretation}"
        )

    document.add_heading("10.7　共用方法限制", level=2)
    limitations = [
        "PC 時序只使用來源資料實際提供且通過共同有效性檢查的時間步；圖中的空白區段保留資料缺口，未推估缺失狀態。",
        "模態排序與空間型態依分析邊界、共同有效遮罩、U₀ 與 E₀ 的定義而定；邊界或尺度改變時應建立新版分析單元並完整重算。",
        "SVD 為線性正交基底，單一物理過程可能分布於相鄰模態，傳播或旋轉訊號亦可能以成對模態表示。",
        "速度異常 RMS 與方向同向性為區域尺度指標；近岸邊界、局地渦旋及來源位置可能使實際漂移方向偏離區域平均向量。",
        "分析期間限定為 2024–2025 年；相近特徵值模態宜以 North sampling error、年度分割或 bootstrap 檢查可分離性與時序穩健性。",
    ]
    for item in limitations:
        document.add_paragraph(item, style="List Bullet")


def add_references_page(document: Document) -> None:
    """在獨立新頁列出正文實際引用的方法學文獻，不附內部成果核對資料。"""

    heading = document.add_heading("參考文獻", level=1)
    # 參考文獻必須獨立成頁；使用標題的分頁前屬性可避免額外空白頁。
    heading.paragraph_format.page_break_before = True
    references = [
        "Bretherton, C. S., Smith, C., & Wallace, J. M. (1992). An intercomparison of methods for finding coupled patterns in climate data. Journal of Climate, 5, 541–560. https://doi.org/10.1175/1520-0442(1992)005<0541:AIOMFF>2.0.CO;2",
        "de Oliveira Júnior, L., Relvas, P., & Garel, E. (2022). Kinematics of surface currents at the northern margin of the Gulf of Cádiz. Ocean Science, 18, 1183–1202. https://doi.org/10.5194/os-18-1183-2022",
        "Golub, G. H., & Van Loan, C. F. (2013). Matrix Computations (4th ed.). Johns Hopkins University Press.",
        "North, G. R., Bell, T. L., Cahalan, R. F., & Moeng, F. J. (1982). Sampling errors in the estimation of empirical orthogonal functions. Monthly Weather Review, 110, 699–706. https://doi.org/10.1175/1520-0493(1982)110<0699:SEITEO>2.0.CO;2",
        "Song, X., Lin, H., Zhan, H., Liu, J., & Cai, S. (2025). Interannual variability of summertime cross-isobath exchanges in the northern South China Sea: ENSO and riverine influences. Ocean Science, 21, 3361–3374. https://doi.org/10.5194/os-21-3361-2025",
        "Volkov, D. L., Lee, S.-K., Landerer, F. W., & Lumpkin, R. (2022). Interannual to decadal sea level variability in the subpolar North Atlantic: the role of propagating signals. Ocean Science, 18, 1741–1762. https://doi.org/10.5194/os-18-1741-2022",
    ]
    for ref in references:
        paragraph = document.add_paragraph(ref)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.28)
        paragraph.paragraph_format.first_line_indent = Inches(-0.28)
        paragraph.paragraph_format.space_after = Pt(5)


def add_synthesis_and_limitations(document: Document, data: Sequence[dict[str, object]], counters: dict[str, int]) -> None:
    """彙整跨區關係、海廢驗證方案、限制與正式結論。"""

    document.add_page_break()
    document.add_heading("11　跨區綜合、海廢輸送意涵與後續驗證", level=1)
    document.add_heading("11.1　六區主導結構", level=2)
    document.add_paragraph(
        "北竿與南竿的前三模態 PC 在共同時段高度相關：Mode 1 約 r=0.94、Mode 2 約 r=0.96，Mode 3 因符號錨定差異"
        "約 r=−0.94。兩區 Mode 1 皆為一致東向成分，顯示馬祖區域可能受共同大尺度強迫控制；但相關不證明海廢由一區"
        "直接傳至另一區，仍須以跨島水道的粒子通過率判定。貢寮與龜山島雖地理相近，模態排序與速度—η 配比不同，"
        "再次說明跨區比較不能只依 Mode 編號。"
    )
    document.add_paragraph(
        "後灣具有約 0.371 m s⁻¹ 的東南向平均向量，遠高於其他五區；在此情況下，海廢淨輸送可能由平均場優先決定，"
        "SVD 模態主要描述其擺動與季節性增減。新竹與馬祖的前兩聯合模態高度濃縮，但其第二模態可含較強海面高度共同"
        "訊號；若研究目標為漂浮物平流，應避免將 η 主導的 EV 直接解釋為強水平輸送。龜山島則呈現相反情況：Mode 2"
        "雖非聯合第一模態，卻具最強且一致的東北向速度回歸場。"
    )

    document.add_heading("11.2　以模態重建進行海廢傳輸驗證", level=2)
    document.add_paragraph(
        "建議以每區平均流與候選模態建立可控制的重建速度場，分別執行單模態、累積模態及完整原始場粒子試驗。對第 K"
        "個重建階數，可令 uᴷ(x,t)=ū(x)+Σₖ₌₁ᴷRᵤ,ₖ(x)PC*ₖ(t)，vᴷ 同理；龜山島另應比較僅 Mode 1 與僅 Mode 2，"
        "後灣則比較平均場、平均場+Mode 1、平均場+Mode 2。"
    )
    validation_rows = [
        ("粒子釋放", "熱點、沿岸格點與區域邊界分層；每小時或每日等量釋放", "避免結果只代表單一初始位置"),
        ("模態情境", "平均場、單一候選模態、前 K 模態、原始完整流場", "量化各模態的增量解釋能力"),
        ("輸出指標", "命中熱點率、離岸／靠岸時間、跨門檻通量、停留時間、FTLE 或連通矩陣", "將空間模態轉成可檢驗輸送量"),
        ("外力敏感度", "加入風壓漂移、Stokes drift、潮汐與擴散係數範圍", "辨識流場模態是否仍為主導"),
        ("穩健性", "年度分割、季節分割、bootstrap、遮罩與尺度敏感度", "量化模態排序與空間型態不確定性"),
    ]
    table = document.add_table(rows=1, cols=3)
    for index, header in enumerate(("驗證面向", "建議設計", "判定意義")):
        table.rows[0].cells[index].text = header
    for values in validation_rows:
        row = table.add_row().cells
        for index, value in enumerate(values):
            row[index].text = value
    set_table_layout(table, [1550, 4300, 3510])
    format_table_text(table, font_size=8.2)
    counters["table"] += 1
    add_table_caption(document, f"表 {counters['table']}　由 SVD 候選模態完成海廢傳輸因果驗證之建議試驗矩陣。")

    document.add_heading("12　限制與不確定性", level=1)
    limitations = [
        "資料可得率約 97.6%，且最大缺口 50 小時；SVD 使用全部可得共同時間，PC 圖保留缺口，未估計缺失時段狀態。",
        "模態依分析區邊界、共同有效遮罩與 u–v／η 尺度定義而變；跨區 EV 不宜視為完全同質母體的統計檢定。",
        "空間模態為正交統計基底，不必然對應單一物理外力；傳播或旋轉訊號可由一對相鄰模態共同表達。",
        "方向同向性與速度 RMS 是區域尺度代理量，不包含海岸碰撞、浮力、風壓漂移、波浪 Stokes drift、潮汐或亂流擴散。",
        "目前僅有 2024–2025 兩年，極端事件與年際穩健性有限；相近特徵值的模態應依 North sampling error 或 bootstrap 檢查可分離性（North et al., 1982）。",
        "候選試驗區尚未等同正式核定分析區；若邊界或熱點定義更新，應建立新版本分析單元並完整重跑，而非覆寫既有成果。",
    ]
    for item in limitations:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("13　結論與建議", level=1)
    conclusions = [
        "面積加權 u–v–η 多變量 SVD 能以單一 PC 同時描述表層速度與海面高度共同變化，且在六區前五模態已解釋 86.07%–98.84% 聯合變異。",
        "北竿、南竿、貢寮與新竹之 Mode 1 同時具最高聯合 EV、較強速度 RMS 與高方向同向性，為首要海廢輸送驗證候選。",
        "龜山島須區分統計主模態與速度主模態：Mode 1 為聯合變異主模態，Mode 2 才是較強且一致的水平速度候選。",
        "後灣的強東南向平均流可能先決定淨輸送；Mode 1 主要控制減弱／反轉，Mode 2 控制同向季節性增強，兩者皆應納入驗證。",
        "最終『主導海廢傳輸模態』應以單模態重建粒子試驗、跨門檻通量、熱點命中率及外力敏感度共同判定，不能只依 EV 排名。",
    ]
    for index, item in enumerate(conclusions, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.add_run(f"{index}. {item}")


def add_references_and_appendix(document: Document, data: Sequence[dict[str, object]], counters: dict[str, int]) -> None:
    """加入方法學文獻與可追溯 provenance 附錄。"""

    document.add_heading("參考文獻", level=1)
    references = [
        "Bretherton, C. S., Smith, C., & Wallace, J. M. (1992). An intercomparison of methods for finding coupled patterns in climate data. Journal of Climate, 5, 541–560. https://doi.org/10.1175/1520-0442(1992)005<0541:AIOMFF>2.0.CO;2",
        "de Oliveira Júnior, L., Relvas, P., & Garel, E. (2022). Kinematics of surface currents at the northern margin of the Gulf of Cádiz. Ocean Science, 18, 1183–1202. https://doi.org/10.5194/os-18-1183-2022",
        "Golub, G. H., & Van Loan, C. F. (2013). Matrix Computations (4th ed.). Johns Hopkins University Press.",
        "North, G. R., Bell, T. L., Cahalan, R. F., & Moeng, F. J. (1982). Sampling errors in the estimation of empirical orthogonal functions. Monthly Weather Review, 110, 699–706. https://doi.org/10.1175/1520-0493(1982)110<0699:SEITEO>2.0.CO;2",
        "Song, X., Lin, H., Zhan, H., Liu, J., & Cai, S. (2025). Interannual variability of summertime cross-isobath exchanges in the northern South China Sea: ENSO and riverine influences. Ocean Science, 21, 3361–3374. https://doi.org/10.5194/os-21-3361-2025",
        "Volkov, D. L., Lee, S.-K., Landerer, F. W., & Lumpkin, R. (2022). Interannual to decadal sea level variability in the subpolar North Atlantic: the role of propagating signals. Ocean Science, 18, 1741–1762. https://doi.org/10.5194/os-18-1741-2022",
    ]
    for ref in references:
        paragraph = document.add_paragraph(ref)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.28)
        paragraph.paragraph_format.first_line_indent = Inches(-0.28)
        paragraph.paragraph_format.space_after = Pt(5)

    document.add_heading("附錄 A　成果可追溯性與重現核對表", level=1)
    headers = ["區域", "run ID", "source metadata SHA-256", "v8 bundle provenance SHA-256"]
    table = document.add_table(rows=1, cols=4)
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for region in data:
        bundle_metadata = region["bundle_metadata"]
        row = table.add_row().cells
        row[0].text = region["spec"].name_zh
        row[1].text = region["spec"].run_id
        row[2].text = bundle_metadata["source_run"]["metadata_sha256"]
        row[3].text = bundle_metadata["bundle_provenance_sha256"]
    set_table_layout(table, [850, 3010, 2750, 2750])
    format_table_text(table, font_size=6.8)
    counters["table"] += 1
    add_table_caption(document, f"表 {counters['table']}　六區 immutable science run 與 academic_report_ready_v8 圖集之內容身分。")

    document.add_paragraph(
        "六個 v8 bundle 均以 Python 3.12.13、NumPy 2.5.1、Matplotlib 3.11.1、FreeType 2.14.3 與 Noto Sans CJK TC"
        " 字型重繪；raster DPI 為 300。空間圖一律採 ``mean_surface_flow_report_with_vector_scale.png`` 或"
        " ``svd_mode_XX_spatial_report_with_vector_scale.png``。完整陣列資料契約包含 ``mean_u/v/eta.npy``、"
        "``regression_u/v/eta.npy``、``pc_standardized.npy``、``time_utc_ns.npy``、``explained_variance.npy``、"
        "``cumulative_explained_variance.npy``、``valid_mask.npy`` 與 ``cell_area_m2.npy``。"
    )


def audit_document(document: Document, output_path: Path) -> None:
    """在儲存前執行結構稽核，避免缺圖、比例尺圖誤用或表格幾何漂移。"""

    # 圖片關係應為 72 張：六區各 1 平均場、1 EV、5 空間模態與 5 PC。
    image_relationships = [rel for rel in document.part.rels.values() if "image" in rel.reltype]
    if len(image_relationships) != 72:
        raise ValueError(f"報告應嵌入 72 張 v8 PNG，實際為 {len(image_relationships)}")
    # 所有表格都應套用固定配置，避免在 Word 中超出 6.5 英寸內容寬。
    for index, table in enumerate(document.tables, start=1):
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        if layout is None or layout.get(qn("w:type")) != "fixed":
            raise ValueError(f"第 {index} 個表格未使用 fixed layout")
    output_path.parent.mkdir(parents=True, exist_ok=True)


def build_report(output_path: Path) -> None:
    """載入六區成果並產生聚焦版學術報告，既有 1.0 版檔案不會被覆寫。"""

    specs_by_key = {spec.key: spec for spec in REGIONS}
    data = [load_region_data(specs_by_key[key]) for key in REPORT_REGION_ORDER]
    document = Document()
    configure_document(document)
    document.core_properties.title = "指定海域流場模態萃取時序變化與主導海廢輸送候選模態報告"
    document.core_properties.subject = "流場模態萃取與時序列變化"
    document.core_properties.author = "OCM-SVD-Analysis"
    document.core_properties.keywords = "SVD, 表層海流, 海洋廢棄物, 多變量分析"
    counters = {"figure": 0, "table": 0}

    add_cover(document)
    add_focused_toc(document)
    add_focused_summary_chapters(document)
    add_method_chapter(document, counters)
    for offset, region in enumerate(data, start=4):
        add_region_chapter(document, region, offset, counters)
    add_independent_conclusions(document, data)
    add_references_page(document)
    enforce_document_font(document)
    audit_document(document, output_path)
    document.save(output_path)


def main() -> int:
    """命令列入口；成功時印出唯一交付檔路徑供後續 render-and-verify 使用。"""

    args = parse_args()
    build_report(args.output.resolve())
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
